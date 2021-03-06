# encoding: utf-8

"""
Test suite for the docx.text module
"""

from __future__ import (
    absolute_import, division, print_function, unicode_literals
)

from docx.enum.text import WD_BREAK, WD_UNDERLINE
from docx.oxml.text import CT_P, CT_R
from docx.text import Paragraph, Run

import pytest

from mock import call, Mock

from .oxml.unitdata.text import (
    a_b, a_bCs, a_br, a_caps, a_cs, a_dstrike, a_p, a_shadow, a_smallCaps,
    a_snapToGrid, a_specVanish, a_strike, a_t, a_u, a_vanish, a_webHidden,
    an_emboss, an_i, an_iCs, an_imprint, an_oMath, a_noProof, an_outline,
    an_r, an_rPr, an_rStyle, an_rtl
)
from .unitutil import class_mock, instance_mock


class DescribeParagraph(object):

    def it_has_a_sequence_of_the_runs_it_contains(self, runs_fixture):
        paragraph, Run_, r_, r_2_, run_, run_2_ = runs_fixture
        runs = paragraph.runs
        assert Run_.mock_calls == [call(r_), call(r_2_)]
        assert runs == [run_, run_2_]

    def it_can_add_a_run_to_itself(self, add_run_fixture):
        paragraph, text, style, expected_xml = add_run_fixture
        run = paragraph.add_run(text, style)
        assert paragraph._p.xml == expected_xml
        assert isinstance(run, Run)
        assert run._r is paragraph._p.r_lst[0]

    def it_knows_its_paragraph_style(self):
        cases = (
            (Mock(name='p_elm', style='foobar'), 'foobar'),
            (Mock(name='p_elm', style=None),     'Normal'),
        )
        for p_elm, expected_style in cases:
            p = Paragraph(p_elm)
            assert p.style == expected_style

    def it_can_set_its_paragraph_style(self):
        cases = (
            ('foobar', 'foobar'),
            ('Normal', None),
        )
        for style, expected_setting in cases:
            p_elm = Mock(name='p_elm')
            p = Paragraph(p_elm)
            p.style = style
            assert p_elm.style == expected_setting

    def it_knows_the_text_it_contains(self, text_prop_fixture):
        p, expected_text = text_prop_fixture
        assert p.text == expected_text

    # fixtures -------------------------------------------------------

    @pytest.fixture(params=[
        (None, None), (None, 'Strong'), ('foobar', None), ('foobar', 'Strong')
    ])
    def add_run_fixture(self, request, paragraph):
        text, style = request.param
        r_bldr = an_r()
        if style:
            r_bldr.with_child(
                an_rPr().with_child(an_rStyle().with_val(style))
            )
        if text:
            r_bldr.with_child(a_t().with_text(text))
        expected_xml = a_p().with_nsdecls().with_child(r_bldr).xml()
        return paragraph, text, style, expected_xml

    # fixture components ---------------------------------------------

    @pytest.fixture
    def p_(self, request, r_, r_2_):
        return instance_mock(request, CT_P, r_lst=(r_, r_2_))

    @pytest.fixture
    def paragraph(self):
        p = a_p().with_nsdecls().element
        return Paragraph(p)

    @pytest.fixture
    def Run_(self, request, runs_):
        run_, run_2_ = runs_
        return class_mock(
            request, 'docx.text.Run', side_effect=[run_, run_2_]
        )

    @pytest.fixture
    def r_(self, request):
        return instance_mock(request, CT_R)

    @pytest.fixture
    def r_2_(self, request):
        return instance_mock(request, CT_R)

    @pytest.fixture
    def runs_(self, request):
        run_ = instance_mock(request, Run, name='run_')
        run_2_ = instance_mock(request, Run, name='run_2_')
        return run_, run_2_

    @pytest.fixture
    def runs_fixture(self, p_, Run_, r_, r_2_, runs_):
        paragraph = Paragraph(p_)
        run_, run_2_ = runs_
        return paragraph, Run_, r_, r_2_, run_, run_2_

    @pytest.fixture
    def text_prop_fixture(self):
        p = (
            a_p().with_nsdecls().with_child(
                an_r().with_child(
                    a_t().with_text('foo'))).with_child(
                an_r().with_child(
                    a_t().with_text(' de bar')))
        ).element
        paragraph = Paragraph(p)
        return paragraph, 'foo de bar'


class DescribeRun(object):

    def it_knows_its_bool_prop_states(self, bool_prop_get_fixture):
        run, prop_name, expected_state = bool_prop_get_fixture
        assert getattr(run, prop_name) == expected_state

    def it_can_change_its_bool_prop_settings(self, bool_prop_set_fixture):
        run, prop_name, value, expected_xml = bool_prop_set_fixture
        setattr(run, prop_name, value)
        assert run._r.xml == expected_xml

    def it_knows_its_character_style(self, style_get_fixture):
        run, expected_style = style_get_fixture
        assert run.style == expected_style

    def it_can_change_its_character_style(self, style_set_fixture):
        run, style, expected_xml = style_set_fixture
        run.style = style
        assert run._r.xml == expected_xml

    def it_knows_its_underline_type(self, underline_get_fixture):
        run, expected_value = underline_get_fixture
        assert run.underline is expected_value

    def it_can_change_its_underline_type(self, underline_set_fixture):
        run, underline, expected_xml = underline_set_fixture
        run.underline = underline
        assert run._r.xml == expected_xml

    def it_raises_on_assign_invalid_underline_type(
            self, underline_raise_fixture):
        run, underline = underline_raise_fixture
        with pytest.raises(ValueError):
            run.underline = underline

    def it_can_add_text(self, add_text_fixture):
        run, text_str, expected_xml, Text_ = add_text_fixture
        _text = run.add_text(text_str)
        assert run._r.xml == expected_xml
        assert _text is Text_.return_value

    def it_can_add_a_break(self, add_break_fixture):
        run, break_type, expected_xml = add_break_fixture
        run.add_break(break_type)
        assert run._r.xml == expected_xml

    def it_knows_the_text_it_contains(self, text_prop_fixture):
        run, expected_text = text_prop_fixture
        assert run.text == expected_text

    # fixtures -------------------------------------------------------

    @pytest.fixture(params=[
        'line', 'page', 'column', 'clr_lt', 'clr_rt', 'clr_all'
    ])
    def add_break_fixture(self, request, run):
        type_, clear, break_type = {
            'line':    (None,           None,    WD_BREAK.LINE),
            'page':    ('page',         None,    WD_BREAK.PAGE),
            'column':  ('column',       None,    WD_BREAK.COLUMN),
            'clr_lt':  ('textWrapping', 'left',  WD_BREAK.LINE_CLEAR_LEFT),
            'clr_rt':  ('textWrapping', 'right', WD_BREAK.LINE_CLEAR_RIGHT),
            'clr_all': ('textWrapping', 'all',   WD_BREAK.LINE_CLEAR_ALL),
        }[request.param]
        # expected_xml -----------------
        br_bldr = a_br()
        if type_ is not None:
            br_bldr.with_type(type_)
        if clear is not None:
            br_bldr.with_clear(clear)
        expected_xml = an_r().with_nsdecls().with_child(br_bldr).xml()
        return run, break_type, expected_xml

    @pytest.fixture(params=['foobar', ' foo bar', 'bar foo '])
    def add_text_fixture(self, request, run, Text_):
        text_str = request.param
        t_bldr = a_t().with_text(text_str)
        if text_str.startswith(' ') or text_str.endswith(' '):
            t_bldr.with_space('preserve')
        expected_xml = an_r().with_nsdecls().with_child(t_bldr).xml()
        return run, text_str, expected_xml, Text_

    @pytest.fixture(params=[
        ('all_caps', True), ('all_caps', False), ('all_caps', None),
        ('bold', True), ('bold', False), ('bold', None),
        ('italic', True), ('italic', False), ('italic', None),
        ('complex_script', True), ('complex_script', False),
        ('complex_script', None),
        ('cs_bold', True), ('cs_bold', False), ('cs_bold', None),
        ('cs_italic', True), ('cs_italic', False), ('cs_italic', None),
        ('double_strike', True), ('double_strike', False),
        ('double_strike', None),
        ('emboss', True), ('emboss', False), ('emboss', None),
        ('hidden', True), ('hidden', False), ('hidden', None),
        ('italic', True), ('italic', False), ('italic', None),
        ('imprint', True), ('imprint', False), ('imprint', None),
        ('math', True), ('math', False), ('math', None),
        ('no_proof', True), ('no_proof', False), ('no_proof', None),
        ('outline', True), ('outline', False), ('outline', None),
        ('rtl', True), ('rtl', False), ('rtl', None),
        ('shadow', True), ('shadow', False), ('shadow', None),
        ('small_caps', True), ('small_caps', False), ('small_caps', None),
        ('snap_to_grid', True), ('snap_to_grid', False),
        ('snap_to_grid', None),
        ('spec_vanish', True), ('spec_vanish', False), ('spec_vanish', None),
        ('strike', True), ('strike', False), ('strike', None),
        ('web_hidden', True), ('web_hidden', False), ('web_hidden', None),
    ])
    def bool_prop_get_fixture(self, request):
        bool_prop_name, expected_state = request.param
        bool_prop_bldr = {
            'all_caps':       a_caps,
            'bold':           a_b,
            'complex_script': a_cs,
            'cs_bold':        a_bCs,
            'cs_italic':      an_iCs,
            'double_strike':  a_dstrike,
            'emboss':         an_emboss,
            'hidden':         a_vanish,
            'italic':         an_i,
            'imprint':        an_imprint,
            'math':           an_oMath,
            'no_proof':       a_noProof,
            'outline':        an_outline,
            'rtl':            an_rtl,
            'shadow':         a_shadow,
            'small_caps':     a_smallCaps,
            'snap_to_grid':   a_snapToGrid,
            'spec_vanish':    a_specVanish,
            'strike':         a_strike,
            'web_hidden':     a_webHidden,
        }[bool_prop_name]
        r_bldr = an_r().with_nsdecls()
        if expected_state is not None:
            child_bldr = bool_prop_bldr()
            if expected_state is False:
                child_bldr.with_val('off')
            rPr_bldr = an_rPr().with_child(child_bldr)
            r_bldr.with_child(rPr_bldr)
        r = r_bldr.element
        run = Run(r)
        return run, bool_prop_name, expected_state

    @pytest.fixture(params=[
        ('all_caps', True), ('all_caps', False), ('all_caps', None),
        ('bold', True), ('bold', False), ('bold', None),
        ('italic', True), ('italic', False), ('italic', None),
        ('complex_script', True), ('complex_script', False),
        ('complex_script', None),
        ('cs_bold', True), ('cs_bold', False), ('cs_bold', None),
        ('cs_italic', True), ('cs_italic', False), ('cs_italic', None),
        ('double_strike', True), ('double_strike', False),
        ('double_strike', None),
        ('emboss', True), ('emboss', False), ('emboss', None),
        ('hidden', True), ('hidden', False), ('hidden', None),
        ('italic', True), ('italic', False), ('italic', None),
        ('imprint', True), ('imprint', False), ('imprint', None),
        ('math', True), ('math', False), ('math', None),
        ('no_proof', True), ('no_proof', False), ('no_proof', None),
        ('outline', True), ('outline', False), ('outline', None),
        ('rtl', True), ('rtl', False), ('rtl', None),
        ('shadow', True), ('shadow', False), ('shadow', None),
        ('small_caps', True), ('small_caps', False), ('small_caps', None),
        ('snap_to_grid', True), ('snap_to_grid', False),
        ('snap_to_grid', None),
        ('spec_vanish', True), ('spec_vanish', False), ('spec_vanish', None),
        ('strike', True), ('strike', False), ('strike', None),
        ('web_hidden', True), ('web_hidden', False), ('web_hidden', None),
    ])
    def bool_prop_set_fixture(self, request):
        bool_prop_name, value = request.param
        bool_prop_bldr = {
            'all_caps':       a_caps,
            'bold':           a_b,
            'complex_script': a_cs,
            'cs_bold':        a_bCs,
            'cs_italic':      an_iCs,
            'double_strike':  a_dstrike,
            'emboss':         an_emboss,
            'hidden':         a_vanish,
            'italic':         an_i,
            'imprint':        an_imprint,
            'math':           an_oMath,
            'no_proof':       a_noProof,
            'outline':        an_outline,
            'rtl':            an_rtl,
            'shadow':         a_shadow,
            'small_caps':     a_smallCaps,
            'snap_to_grid':   a_snapToGrid,
            'spec_vanish':    a_specVanish,
            'strike':         a_strike,
            'web_hidden':     a_webHidden,
        }[bool_prop_name]
        # run --------------------------
        r = an_r().with_nsdecls().element
        run = Run(r)
        # expected_xml -----------------
        rPr_bldr = an_rPr()
        if value is not None:
            child_bldr = bool_prop_bldr()
            if value is False:
                child_bldr.with_val(0)
            rPr_bldr.with_child(child_bldr)
        expected_xml = an_r().with_nsdecls().with_child(rPr_bldr).xml()
        return run, bool_prop_name, value, expected_xml

    @pytest.fixture(params=['Foobar', None])
    def style_get_fixture(self, request):
        style = request.param
        r = self.r_bldr_with_style(style).element
        run = Run(r)
        return run, style

    @pytest.fixture(params=[
        (None, None),
        (None, 'Foobar'),
        ('Foobar', None),
        ('Foobar', 'Foobar'),
        ('Foobar', 'Barfoo'),
    ])
    def style_set_fixture(self, request):
        before_style, after_style = request.param
        r = self.r_bldr_with_style(before_style).element
        run = Run(r)
        expected_xml = self.r_bldr_with_style(after_style).xml()
        return run, after_style, expected_xml

    @pytest.fixture
    def text_prop_fixture(self, Text_):
        r = (
            an_r().with_nsdecls().with_child(
                a_t().with_text('foo')).with_child(
                a_t().with_text('bar'))
        ).element
        run = Run(r)
        return run, 'foobar'

    @pytest.fixture(params=[
        (None,     None),
        ('single', True),
        ('none',   False),
        ('double', WD_UNDERLINE.DOUBLE),
    ])
    def underline_get_fixture(self, request):
        underline_type, expected_prop_value = request.param
        r = self.r_bldr_with_underline(underline_type).element
        run = Run(r)
        return run, expected_prop_value

    @pytest.fixture(params=['foobar', 42, 'single'])
    def underline_raise_fixture(self, request):
        underline = request.param
        r = self.r_bldr_with_underline(None).element
        run = Run(r)
        return run, underline

    @pytest.fixture(params=[
        (None,     True,                'single'),
        (None,     False,               'none'),
        (None,     None,                None),
        (None,     WD_UNDERLINE.SINGLE, 'single'),
        (None,     WD_UNDERLINE.WAVY,   'wave'),
        ('single', True,                'single'),
        ('single', False,               'none'),
        ('single', None,                None),
        ('single', WD_UNDERLINE.SINGLE, 'single'),
        ('single', WD_UNDERLINE.DOTTED, 'dotted'),
    ])
    def underline_set_fixture(self, request):
        before_val, underline, expected_val = request.param
        r = self.r_bldr_with_underline(before_val).element
        run = Run(r)
        expected_xml = self.r_bldr_with_underline(expected_val).xml()
        return run, underline, expected_xml

    # fixture components ---------------------------------------------

    @pytest.fixture
    def run(self):
        r = an_r().with_nsdecls().element
        return Run(r)

    def r_bldr_with_style(self, style):
        rPr_bldr = an_rPr()
        if style is not None:
            rPr_bldr.with_child(an_rStyle().with_val(style))
        r_bldr = an_r().with_nsdecls().with_child(rPr_bldr)
        return r_bldr

    def r_bldr_with_underline(self, underline_type):
        rPr_bldr = an_rPr()
        if underline_type is not None:
            rPr_bldr.with_child(a_u().with_val(underline_type))
        r_bldr = an_r().with_nsdecls().with_child(rPr_bldr)
        return r_bldr

    @pytest.fixture
    def Text_(self, request):
        return class_mock(request, 'docx.text.Text')
